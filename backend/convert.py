#!/usr/bin/env python3
"""
convert.py — 將 data/other/ 來源檔批次轉換為設備警報 JSON
輸出到 data/{device_model}.json

用法: python3 backend/convert.py
注意: TFM001 為掃描版 PDF，無法自動解析，已略過需手動處理
"""

import json
import re
from pathlib import Path

BASE = Path(__file__).resolve().parent.parent
SRC  = BASE / "data" / "other"
DEST = BASE / "data"

# ── 來源檔 → device_model 對應 ───────────────────────────────────────────────

FILE_MAP = {
    # Excel
    "ACM001警報清單.xlsx":   "ACM001",
    "ACM002警報.xlsx":       "ACM002",
    "SDM003.xlsx":           "SDDM003",
    "SDM004.xlsx":           "SDDM004",
    # Word
    "ACP001  警報清單.docx": "ACP001",
    "ACP002 警報列表.docx":  "ACP002",
    "PILM003警報列表.docx":  "PILM003",
    "PILM004 警報列表.docx": "PILM004",
    "TFM002警報列表.docx":   "TFM002",
    # PDF（義大利文/英文雙語，字元重複版）
    "(DTS002_SL-A210)List of indications_001_90_01113_ITA_ENG.pdf": "DTS002-SL-A210",
    "(DTS002_BL-A420)List of indications_001_90_01113_ITA_ENG.pdf": "DTS002-BL-A420",
    "(DTS003_SL-A210)List of indications_001_90_01113_ITA_ENG.pdf": "DTS003-SL-A210",
    "(DTS003_BL-A420)List of indications_001_90_01113_ITA_ENG.pdf": "DTS003-BL-A420",
    # TFM001 — 掃描版 PDF，略過
}

# ── 通用 ──────────────────────────────────────────────────────────────────────

def make_alarm(code, model, desc="", cause="", solution="", keywords=None):
    return {
        "code":         str(code).zfill(4),
        "device_model": model,
        "severity":     "",
        "description":  desc.strip(),
        "cause":        cause.strip(),
        "solution":     solution.strip(),
        "keywords":     keywords or [],
    }

# ── Excel ─────────────────────────────────────────────────────────────────────

def parse_excel(path: Path, model: str) -> list:
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True)

    # 找第一個有資料的 sheet
    ws = None
    for name in wb.sheetnames:
        sheet = wb[name]
        if sheet.max_row > 1:
            ws = sheet
            break
    if not ws:
        return []

    alarms = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        code_val = row[0]
        if code_val is None:
            continue
        try:
            # 處理 "0001-"、"0002 -" 等尾綴 dash 格式
            code_str = re.sub(r'[\s\-]+$', '', str(code_val))
            code = str(int(code_str)).zfill(4)
        except (ValueError, TypeError):
            continue

        eng = str(row[1] or "").strip()
        chn = str(row[2] or "").strip()
        desc = f"{eng} {chn}".strip() if chn else eng
        if not desc:
            continue

        alarms.append(make_alarm(code, model, desc=desc))

    return alarms

# ── Word ──────────────────────────────────────────────────────────────────────

CODE_RE = re.compile(r'^(\d{1,4})\s*[-–]\s*([A-Za-z一-鿿].+)$')

SOLUTION_PREFIXES = re.compile(
    r'^(Check|Verify|Remove|Replace|Restore|Ensure|Reset|Adjust|Clean|Inspect|'
    r'Press|Contact|Do not|確認|檢查|更換|清理|移除|重新|啟動|關閉|按下)',
    re.IGNORECASE,
)

def parse_word(path: Path, model: str) -> list:
    import docx
    doc = docx.Document(path)

    if doc.tables:
        return _parse_word_table(doc, model)
    return _parse_word_paragraphs(doc, model)

def _parse_word_paragraphs(doc, model: str) -> list:
    alarms = []
    current = None
    cause_lines = []
    solution_lines = []
    in_solution = False

    def _flush():
        if current:
            current["cause"]    = " ".join(cause_lines).strip()
            current["solution"] = "\n".join(solution_lines).strip()
            alarms.append(current)

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        m = CODE_RE.match(text)
        if m:
            _flush()
            code = str(int(m.group(1))).zfill(4)
            current = make_alarm(code, model, desc=m.group(2).strip())
            cause_lines    = []
            solution_lines = []
            in_solution    = False
        elif current:
            if SOLUTION_PREFIXES.match(text) or text.startswith(('-', '–', '•')):
                in_solution = True
                solution_lines.append(text.lstrip('-–•').strip())
            elif not in_solution:
                cause_lines.append(text)
            else:
                solution_lines.append(text)

    _flush()
    return alarms

def _parse_word_table(doc, model: str) -> list:
    alarms = []
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if not cells or not cells[0]:
                continue
            try:
                code = str(int(cells[0])).zfill(4)
            except (ValueError, TypeError):
                continue
            desc = cells[1] if len(cells) > 1 else ""
            alarms.append(make_alarm(code, model, desc=desc))
    return alarms

# ── PDF（DTS 系列，字元重複版雙語 PDF）────────────────────────────────────────

# 義大利文常見開頭，用來辨別義大利段落
ITALIAN_RE = re.compile(
    r'^(il |la |lo |le |che |non |nella |del |della |per |con |una |uno |gli |da |al |è |si |verificare|togliere|premere)',
    re.IGNORECASE,
)

def _dedup(s: str) -> str:
    """'PPRROODDUUCCTT LLOOAADDIINNGG' → 'PRODUCT LOADING'（逐詞去重，避免空格落在奇數位置被跳過）"""
    return ' '.join(''.join(w[i] for i in range(0, len(w), 2)) for w in s.split())

# 每個字元都重複兩次的 code 行，例如: "00000088 -- PPRROODDUUCCTT..."
DOUBLED_CODE_RE = re.compile(r'^((?:\d\d){1,5})\s*--+\s+(.+)$')

HEADER_RE = re.compile(r'^(Elenco segnalazioni|List of indications|\d+\s*/\s*\d+|N\d{6,})')

def _is_italian_block(lines: list) -> bool:
    for l in lines[:3]:
        if l.strip() and ITALIAN_RE.match(l.strip()):
            return True
    return False

def parse_pdf(path: Path, model: str) -> list:
    import pdfplumber

    lines = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                lines.extend(text.splitlines())

    # 找所有 doubled code 行的位置
    entries = []  # (line_idx, code, title)
    for i, line in enumerate(lines):
        m = DOUBLED_CODE_RE.match(line.strip())
        if m:
            code  = _dedup(m.group(1)).zfill(4)
            title = _dedup(m.group(2))
            entries.append((i, code, title))

    alarms_dict = {}
    for idx, (line_idx, code, title) in enumerate(entries):
        end_idx = entries[idx + 1][0] if idx + 1 < len(entries) else len(lines)

        body = [l.strip() for l in lines[line_idx + 1:end_idx]
                if l.strip() and not HEADER_RE.match(l.strip())]

        # 跳過義大利文段落
        if _is_italian_block(body):
            continue

        cause_lines, sol_lines = [], []
        in_sol = False
        for l in body:
            if SOLUTION_PREFIXES.match(l):
                in_sol = True
            (sol_lines if in_sol else cause_lines).append(l)

        # 同一 code 可能英文出現一次，用 dict 去重
        if code not in alarms_dict:
            alarms_dict[code] = make_alarm(
                code, model,
                desc     = title,
                cause    = " ".join(cause_lines),
                solution = "\n".join(sol_lines),
            )

    return list(alarms_dict.values())

# ── 掃描版 PDF（TFM001，使用 OCR）────────────────────────────────────────────

TESSERACT_CMD = "/opt/homebrew/bin/tesseract"

def parse_scanned_pdf_folder(folder: Path, model: str) -> list:
    """將資料夾內所有 PDF 依頁 OCR，合併後用 Word 段落解析器處理。
    自動偵測頁面方向：若正常方向未偵測到 alarm code，則旋轉 180° 重試。"""
    import pytesseract
    import pdfplumber
    from PIL import Image

    pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD
    _has_code = re.compile(r'^\d{1,4}\s*[-–]\s*[A-Za-z一-鿿]')

    def _ocr(img):
        return pytesseract.image_to_string(img, lang="chi_tra+eng")

    all_paragraphs = []
    for pdf_path in sorted(folder.glob("*.pdf")):
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                img  = page.to_image(resolution=200).original
                text = _ocr(img)
                # 若沒找到任何 code 行，旋轉 180° 重試
                if not any(_has_code.match(l.strip()) for l in text.splitlines()):
                    text = _ocr(img.rotate(180))
                all_paragraphs.extend(text.splitlines())

    # 用假 docx 結構重用段落解析器
    class _FakePara:
        def __init__(self, text): self.text = text

    class _FakeDoc:
        def __init__(self, lines):
            self.paragraphs = [_FakePara(l) for l in lines]
            self.tables = []

    return _parse_word_paragraphs(_FakeDoc(all_paragraphs), model)

# ── 主程式 ────────────────────────────────────────────────────────────────────

def main():
    total_alarms = 0
    success = 0

    for filename, model in FILE_MAP.items():
        path = SRC / filename
        if not path.exists():
            print(f"⚠  找不到: {filename}")
            continue

        suffix = path.suffix.lower()
        try:
            if suffix == ".xlsx":
                alarms = parse_excel(path, model)
            elif suffix == ".docx":
                alarms = parse_word(path, model)
            elif suffix == ".pdf":
                alarms = parse_pdf(path, model)
            else:
                print(f"⚠  不支援格式: {filename}")
                continue
        except Exception as e:
            print(f"✗  {model} ({filename}): {e}")
            continue

        out = DEST / f"{model}.json"
        out.write_text(json.dumps(alarms, ensure_ascii=False, indent=2), encoding="utf-8")
        print(f"✓  {model:20s} {len(alarms):>3} 筆")
        total_alarms += len(alarms)
        success += 1

    # TFM001 掃描版 PDF（資料夾）
    tfm001_folder = SRC / "TFM001"
    if tfm001_folder.is_dir():
        try:
            alarms = parse_scanned_pdf_folder(tfm001_folder, "TFM001")
            out = DEST / "TFM001.json"
            out.write_text(json.dumps(alarms, ensure_ascii=False, indent=2), encoding="utf-8")
            print(f"✓  {'TFM001':20s} {len(alarms):>3} 筆  (OCR)")
            total_alarms += len(alarms)
            success += 1
        except Exception as e:
            print(f"✗  TFM001 OCR 失敗: {e}")
    else:
        print("⚠  TFM001 資料夾找不到")

    print(f"\n完成 {success}/{len(FILE_MAP) + 1} 個設備，共 {total_alarms} 筆警報")


if __name__ == "__main__":
    main()
