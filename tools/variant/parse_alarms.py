#!/usr/bin/env python3
"""原廠警報文件 → 標準 JSON（含 variant 欄位）

支援 .xlsx/.xlsm/.csv、.docx、.pdf（僅文字型；掃描件需先光柵化走 AI 管線）。

設計原則（見 EXTRACTION_SPEC.md）：
  * 不寫通用解析器。共用層（去重、正規化、品質檢查、輸出格式）寫死，
    來源層（怎麼從檔案讀出原始列）每種格式各自一個 reader。
  * 規則優先，AI 備援。結構規整的來源（Excel 表格）純規則就解得乾淨，
    動用 AI 反而更差（幻覺風險、要校對、花錢、要等）。
  * 解析結果一律需人工確認才進資料庫。這是製藥設備的處置指示，
    現場人員會照做。

用法：
    python parse_alarms.py -i FILL203.xlsx -m FILL203 -o out.json
    python parse_alarms.py -i 手冊.docx -m ACM002 --action-to local_solution
    python parse_alarms.py -i 手冊.pdf -m ACM002            # 掃描件會提示改走光柵化
    python parse_alarms.py -i FILL203.xlsx -m FILL203 --dry-run
"""
from __future__ import annotations

import argparse
import json
import pathlib
import re
import sys
from collections import defaultdict
from typing import Iterable

# read_tabular()/_detect_columns() 現在是 backend/alarm_ingest/detect.py
# 的共用實作（後台批次匯入 UI 的 inspect 端點也用同一份）——分岔的代價
# 是同一份原廠檔案在 CLI 與後台跑出不同的欄位對應，比 normalize_variant
# 那種各自一份的分岔風險高很多，所以這裡改用 import 而非複製。CLI 依賴
# backend/ 而非反過來，因為 backend/ 是實際部署的那個。
sys.path.insert(0, str(pathlib.Path(__file__).resolve().parent.parent.parent / "backend"))
from alarm_ingest.detect import (  # noqa: E402
    read_tabular,
    list_sheets,
    detect_columns as _detect_columns,
    grid_to_rows as _grid_to_rows,
    split_code,
)

# ── 常數 ────────────────────────────────────────────────────────────

OUTPUT_FIELDS = ["code", "device_model", "variant", "severity",
                 "description", "cause", "solution", "keywords"]

NA_VALUES = {"#N/A", "N/A", "NA", "-", "—", "", "None", "nan"}

# 元件編號：B42.0 / F143.5 / QF160.1 / K120.6
COMPONENT_RE = re.compile(r"\b[A-Z]{1,3}\d+\.\d+\b")
# 疑似掉了小數點：B420（OCR 常見）
SUSPECT_COMPONENT_RE = re.compile(r"\b[A-Z]{1,3}\d{3,}\b")

# 音譯亂碼常用字。既有資料實測命中：
#   MANCANZA TENSIONE DI RETE → 曼坎薩緊張迪雷特
#   ARRESTO MANUALE → 逮捕手冊 / REPHASE MACHINE → 複相機
# 這類錯誤不是空白，是通順但錯誤的內容，只有規則抓得到。
PHONETIC_CHARS = set("薩坎納曼蒂庫西迪雷特阿比埃")

# 簡體字混入繁體語境的標記字。只標記，不自動轉換——一簡對多繁的字
# （发→發/髮）自動轉會出錯，而處置說明錯一個字可能改變操作意思
# （TFM001 0139 實測案例：AI 輸出「冷却板」，原圖是「冷卻板」）。
SIMPLIFIED_ONLY_CHARS = set("却开关闭电压设备温阀门检查连接线圈显示屏")


# ── 共用工具 ────────────────────────────────────────────────────────

def clean(v) -> str:
    """壓掉換行與連續空白，#N/A 類值視為空。"""
    if v is None:
        return ""
    s = " ".join(str(v).split())
    return "" if s in NA_VALUES else s


def normalize_variant(s: str) -> str:
    """variant 進主鍵，任何字元差異都是不同的警報。做不影響顯示的正規化，
    避免來源微幅變動（空白、標點變體）產生重複列並讓 local_solution 脫鉤。

    實際會發生的情境：原廠改版把 hyphen 換成 en dash、Excel 重新匯出多了
    一個空白 → 主鍵不同 → INSERT 新列而非 UPDATE → 現場查到兩筆幾乎一樣的
    變體，而人工累積的 local_solution 留在那筆沒人會選的舊列上。

    刻意不做大小寫轉換——原廠標題大小寫穩定，轉了反而讓顯示變醜。"""
    s = " ".join(s.split())
    s = s.replace("\u2013", "-").replace("\u2014", "-")   # en/em dash → hyphen
    s = s.replace("\uff08", "(").replace("\uff09", ")")   # 全形括號
    s = s.replace("\uff0f", "/")
    return s.strip()


def decide_variant_mode(rows: list[dict], mode: str) -> tuple[bool, str]:
    """回傳 (是否啟用 variant, 判定理由)。mode 為 auto 時依代碼是否唯一判斷。

    純函式，不做 I/O——與後台批次匯入 API 共用同一份契約（各自維護一份
    複本，見 backend/alarm_ingest/quality.py 開頭說明；Variant/ 不進版控，
    無法用 import 共用，靠 backend/tests/test_alarm_ingest.py 的
    DECIDE_VARIANT_CASES 契約測試釘住兩邊同步），呼叫端自行決定要不要
    印出理由。

    auto 有誤判風險：來源若只是全部警報的子集（FILL203 的批次報告就是，
    Problem 分頁自己寫明「部分 alarm 未列於清單」），可能剛好每個代碼只
    出現一次而被判成 never，然後匯入時跟既有的多變體資料撞在一起。
    所以 --variant-mode 設為必填，auto 要明確指定才啟用。"""
    if mode == "always":
        return True, "手動指定啟用 variant"
    if mode == "never":
        return False, "手動指定不啟用 variant"
    codes = [r["code"] for r in rows]
    unique = len(codes) == len(set(codes))
    if unique:
        reason = f"本來源 {len(set(codes))} 個代碼全部唯一，variant 設為空字串"
    else:
        reason = f"偵測到代碼重複（{len(codes)} 列 / {len(set(codes))} 個代碼），啟用 variant"
    return not unique, reason


# read_tabular()/_detect_columns()/split_code() 是 backend/alarm_ingest/
# detect.py 的共用實作（見檔案頂部 import）——後台批次匯入 UI 的 inspect
# 端點用同一份，避免同一份原廠檔案在 CLI 與後台跑出不同的欄位對應。


# ── 來源層：Word ────────────────────────────────────────────────────

def read_docx(path: pathlib.Path) -> list[dict]:
    """讀 .docx。先試表格，沒有表格才走段落模式。

    段落模式對應原廠手冊的常見結構：
        0024 - Forming Panel Heating Units Automatic  成型加熱站點位異常
        The message indicates that a circuit breaker...      ← cause
        - Check that the temperature of the units...         ← solution（原廠）
        檢查加熱板溫度是否過高；檢查Sensor B42.0是否鬆脫      ← 本廠做法

    ⚠️ 未以真實 .docx 驗證過（手上只有掃描 PDF 版本）。第一次使用時
    務必先跑 --dry-run 並人工核對前 10 筆。
    """
    import docx
    doc = docx.Document(str(path))
    rows: list[dict] = []

    for t_i, table in enumerate(doc.tables):
        grid = [[c.text for c in row.cells] for row in table.rows]
        cols = _detect_columns(grid)
        if cols is None:
            continue
        rows.extend(_grid_to_rows(grid, cols, source=f"{path.name}#table{t_i}"))
    if rows:
        print(f"  表格模式：{len(rows)} 列", file=sys.stderr)
        return rows

    rows = _parse_manual_paragraphs([p.text for p in doc.paragraphs], path.name)
    print(f"  段落模式：{len(rows)} 列", file=sys.stderr)
    return rows


def _parse_manual_paragraphs(lines: Iterable[str], source: str) -> list[dict]:
    """手冊式段落結構的共用解析（Word 段落與文字型 PDF 都用這支）。"""
    rows: list[dict] = []
    cur: dict | None = None

    def flush():
        if cur and cur["code"]:
            cur["cause"] = " ".join(cur.pop("_cause")).strip()
            zh = " ".join(cur.pop("_zh")).strip()
            en = "\n".join(cur.pop("_en")).strip()
            # 中文段落是本廠自行加註（含 B42.0/F43.1 這類本廠元件編號），
            # 英文 bullet 才是原廠處置。兩者語意不同，不可混為一談。
            cur["action"] = zh or en
            cur["_vendor_solution"] = en
            rows.append(cur)

    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        m = re.match(r"^(\d{3,6})\s*[-–—]\s*(.+)$", line)
        if m:
            flush()
            cur = {"code": m.group(1), "variant": clean(m.group(2)),
                   "_cause": [], "_en": [], "_zh": [], "_source": source}
            continue
        if cur is None:
            continue
        if line.startswith(("-", "•", "‧")):
            cur["_en"].append(line)
        elif re.search(r"[\u4e00-\u9fff]", line):
            cur["_zh"].append(line)
        else:
            cur["_cause"].append(line)
    flush()
    return rows


# ── 來源層：PDF ─────────────────────────────────────────────────────

def read_pdf(path: pathlib.Path) -> list[dict]:
    """讀 .pdf。無文字層（掃描件）時中止並提示改走光柵化 + AI 管線。"""
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    total_chars = sum(len(p.extract_text() or "") for p in reader.pages)

    if total_chars < 50:
        rotations = {p.get("/Rotate", 0) for p in reader.pages}
        sys.exit(
            f"\n✗ {path.name} 是掃描件（{len(reader.pages)} 頁、文字層 {total_chars} 字元），"
            f"規則式解析無法處理。\n"
            f"  頁面旋轉：{sorted(rotations)}"
            f"{'  ← 不一致，務必光柵化整頁而非抽內嵌影像' if len(rotations) > 1 else ''}\n\n"
            f"  請改走：python parse_alarms.py --rasterize -i {path.name} -o pages/\n"
            f"  再把 PNG 逐頁送 tools/variant/EXTRACTION_SPEC.md 第 4 節的 Gemini prompt。\n"
        )

    import pdfplumber
    lines: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            lines.extend((page.extract_text() or "").splitlines())
    rows = _parse_manual_paragraphs(lines, path.name)
    print(f"  文字型 PDF：{len(rows)} 列", file=sys.stderr)
    return rows


def rasterize(path: pathlib.Path, outdir: pathlib.Path, scale: float = 3.0) -> None:
    """掃描 PDF 光柵化成 PNG，供視覺辨識使用。

    必須光柵化整頁而非抽取內嵌影像：頁面 /Rotate 可能是 0/180/90 三種，
    render() 會套用旋轉，page.images 不會（會得到顛倒的影像，OCR 全是垃圾）。
    """
    import pypdfium2 as pdfium
    outdir.mkdir(parents=True, exist_ok=True)
    pdf = pdfium.PdfDocument(str(path))
    for i, page in enumerate(pdf, start=1):
        img = page.render(scale=scale).to_pil()
        if img.width > img.height:      # 橫向頁轉正，否則辨識率明顯下降
            img = img.rotate(-90, expand=True)
            print(f"  第 {i} 頁為橫向，已轉正")
        out = outdir / f"page_{i:02d}.png"
        img.save(out)
        print(f"  第 {i} 頁 → {out}  {img.size}")


# ── 共用層：去重 ────────────────────────────────────────────────────

def dedup(rows: list[dict]) -> tuple[list[dict], int]:
    """依 (code, variant) 去重，同鍵取內容最完整者。回傳 (結果, 衝突組數)。"""
    best: dict[tuple[str, str], dict] = {}
    seen: dict[tuple[str, str], set] = defaultdict(set)
    for r in rows:
        k = (r["code"], r["variant"])
        seen[k].add((r.get("cause", ""), r.get("action", "")))
        score = len(r.get("cause", "")) + len(r.get("action", ""))
        if k not in best or score > len(best[k].get("cause", "")) + len(best[k].get("action", "")):
            best[k] = r
    conflicts = sum(1 for v in seen.values() if len(v) > 1)
    return sorted(best.values(), key=lambda x: (x["code"], x["variant"])), conflicts


# ── 共用層：品質檢查 ────────────────────────────────────────────────

def quality_check(rows: list[dict]) -> list[tuple[str, str, str]]:
    """回傳 [(識別, 等級, 訊息)]。等級：ERROR / WARN / INFO。"""
    issues: list[tuple[str, str, str]] = []
    codes = [r["code"] for r in rows]

    # 只有當同一個字母前綴在文件別處確實出現過帶小數點的形式，才把無小數點
    # 的版本視為疑似 OCR 掉點。否則像 FILL203 的閥號 V46004 會被大量誤報
    # ——那類編號本來就沒有小數點（實測 36 項警告中 27 項是這種誤判）。
    all_text = " ".join(f'{r.get("cause","")} {r.get("action","")}' for r in rows)
    decimal_prefixes = {m[:re.search(r"\d", m).start()]
                        for m in COMPONENT_RE.findall(all_text)}

    for r in rows:
        key = f'{r["code"]}/{r["variant"][:24]}' if r["variant"] else r["code"]
        blob = f'{r.get("cause","")} {r.get("action","")}'

        for m in SUSPECT_COMPONENT_RE.findall(blob):
            prefix = m[:re.search(r"\d", m).start()]
            if prefix in decimal_prefixes:
                issues.append((key, "WARN", f"元件編號疑似缺小數點：{m}"))

        if "%1" in r["variant"] and "%1" not in blob:
            issues.append((key, "WARN", "標題含 %1 但內文未出現，可能被誤填"))

        zh = r["variant"]
        if zh and sum(c in PHONETIC_CHARS for c in zh) >= 3:
            issues.append((key, "WARN", f"疑似音譯亂碼：{zh[:30]}"))

        zh_all = f'{zh} {blob}'
        simplified_hits = sorted(set(zh_all) & SIMPLIFIED_ONLY_CHARS)
        if simplified_hits:
            issues.append((key, "WARN", f"簡繁混用（不自動轉換，人工確認）：{''.join(simplified_hits)}"))

        if not r.get("cause") and not r.get("action"):
            issues.append((key, "INFO", "無原因也無處置"))

    # 跳號門檻用相對值。寫死絕對值會在不同原廠體系間失準：義大利系四位數
    # 連續編號，Bosch/Syntegon 五位數依子系統分區段（12xxx/30xxx/92xxx），
    # 且批次報告本身就是全部警報的子集，大間隔是常態（實測誤報 8 項）。
    nums = sorted(int(c) for c in set(codes) if c.isdigit())
    gaps = [b - a for a, b in zip(nums, nums[1:])]
    if len(gaps) >= 5:
        median = sorted(gaps)[len(gaps) // 2]
        threshold = max(100, median * 20)
        for a, b in zip(nums, nums[1:]):
            if b - a > threshold:
                issues.append((f"{a}→{b}", "WARN",
                               f"代碼跳號 {b - a}（中位數 {median}），確認是否漏抽整段"))

    return issues


# ── 共用層：輸出 ────────────────────────────────────────────────────

def to_output(rows: list[dict], device_model: str, action_to: str,
              use_variant: bool, desc_format: str) -> list[dict]:
    """組出資料庫格式。

    description 預設無前綴（plain），與既有 1759 筆一致——前台代碼本來就
    另外顯示，加前綴是重複資訊，且在多變體清單上會讓 16 個選項全部以
    同一組數字開頭，使用者要跳過那段才看得到區別文字。搜尋不受影響：
    list_alarms 的比對字串已經把 code 併進去了。"""
    out = []
    for r in rows:
        title = r["variant"]
        desc = f'{r["code"]} {title}'.strip() if desc_format == "with-code" else title
        rec = {
            "code": r["code"],
            "device_model": device_model,
            "variant": normalize_variant(title) if use_variant else "",
            "severity": "",              # 既有資料 100% 空白，不使用
            "description": desc,
            "cause": r.get("cause", ""),
            "solution": "",
            "local_solution": "",
            "keywords": [],
        }
        rec[action_to] = r.get("action", "")
        if r.get("_vendor_solution") and action_to == "local_solution":
            rec["solution"] = r["_vendor_solution"]
        if action_to == "solution":
            rec.pop("local_solution")
        out.append(rec)
    return out


def compare_against(new_rows: list[dict], path: pathlib.Path) -> list[str]:
    """跟既有資料比對，找出「同代碼但 variant 措辭不同」的疑似變動。

    第二次匯入同一台機種時才用得到，但那正是最容易出事、也最不會有人
    警覺的時候：措辭一變主鍵就變，會 INSERT 新列而非 UPDATE，人工累積的
    local_solution 留在舊列上。

    既有資料的取得（維持本工具不碰資料庫）：
        curl -b cookie.txt "$BASE/api/alarms?dept=mf4d" > existing.json
    """
    import difflib
    existing = json.loads(path.read_text(encoding="utf-8"))
    by_code: dict[str, list[dict]] = defaultdict(list)
    for e in existing:
        by_code[str(e.get("code", ""))].append(e)

    warnings = []
    for r in new_rows:
        olds = by_code.get(r["code"])
        if not olds:
            continue
        nv = r["variant"]
        if any(normalize_variant(o.get("variant", "")) == normalize_variant(nv) for o in olds):
            continue  # 完全相符，不是變動
        for o in olds:
            ov = o.get("variant", "")
            ratio = difflib.SequenceMatcher(None, normalize_variant(ov),
                                            normalize_variant(nv)).ratio()
            if ratio >= 0.85:
                has_local = bool((o.get("local_solution") or "").strip())
                warnings.append(
                    f'{r["code"]} 疑似措辭變動（相似度 {ratio:.2f}）'
                    f'{"，且既有那筆已有 local_solution" if has_local else ""}\n'
                    f'       既有：{ov[:70]}\n'
                    f'       新的：{nv[:70]}'
                )
                break
    return warnings


# ── main ────────────────────────────────────────────────────────────

def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("-i", "--input", required=True, type=pathlib.Path)
    ap.add_argument("-m", "--device-model", help="寫進 device_model 欄位（--rasterize 時不需要）")
    ap.add_argument("-o", "--output", type=pathlib.Path)
    ap.add_argument("--action-to", choices=["solution", "local_solution"], default="solution",
                    help="來源的處置欄要寫進哪個欄位。Excel 批次報告的 Action 欄是"
                         "本廠自行撰寫的，嚴格說應為 local_solution（預設 solution "
                         "是為了相容 local_solution 欄位尚未上線的情況）")
    ap.add_argument("--variant-mode", choices=["auto", "always", "never"], default=None,
                    help="必填。always=標題進 variant（一碼多變體的來源，如 Bosch/"
                         "Syntegon 體系）；never=variant 留空（代碼唯一的來源，如既有"
                         "義大利系機種）；auto=依代碼是否唯一自動判斷（子集來源可能誤判）")
    ap.add_argument("--description-format", choices=["plain", "with-code"], default="plain",
                    help="plain（預設）＝無代碼前綴，與既有 1759 筆一致")
    ap.add_argument("--against", type=pathlib.Path,
                    help="既有資料 JSON，用於偵測 variant 措辭變動。取得方式："
                         "curl -b cookie.txt \"$BASE/api/alarms?dept=mf4d\" > existing.json")
    ap.add_argument("--dry-run", action="store_true", help="只印統計與品質檢查，不寫檔")
    ap.add_argument("--rasterize", action="store_true", help="掃描 PDF 轉 PNG（-o 為輸出目錄）")
    ap.add_argument("--sheet", help="只讀指定分頁（.xlsx/.xlsm 適用）。不指定則自動掃描"
                    "全部分頁——多分頁檔案若只有部分分頁是要匯入的資料，用這個參數")
    ap.add_argument("--sheet-list", action="store_true",
                    help="只列出分頁名稱與是否偵測到警報欄位，不解析內容，用來"
                         "確認多分頁檔案該用哪個 --sheet")
    args = ap.parse_args()

    if not args.input.exists():
        sys.exit(f"✗ 找不到檔案：{args.input}")

    if args.sheet_list:
        for name, detected in list_sheets(args.input):
            print(f"  {'✓' if detected else ' '} {name!r}")
        return

    if args.rasterize:
        rasterize(args.input, args.output or pathlib.Path("pages"))
        return

    if not args.device_model:
        sys.exit("✗ 需要 --device-model")
    if args.variant_mode is None:
        sys.exit("✗ 需要 --variant-mode {always|never|auto}\n"
                 "  always：一碼多變體的來源（Bosch/Syntegon 體系）\n"
                 "  never ：代碼唯一的來源（既有義大利系機種）\n"
                 "  auto  ：自動判斷（子集來源可能誤判，會印出判定結果）\n"
                 "  設為必填而非預設 auto，是因為誤判的後果是產生重複列並讓\n"
                 "  既有的 local_solution 脫鉤，而那不會有任何錯誤訊息。")

    if args.sheet and args.input.suffix.lower() not in (".xlsx", ".xlsm"):
        sys.exit("✗ --sheet 只適用於 .xlsx/.xlsm")

    suffix = args.input.suffix.lower()
    print(f"讀取 {args.input.name}（{suffix}）", file=sys.stderr)
    if suffix in (".xlsx", ".xlsm", ".csv"):
        raw = read_tabular(args.input, sheet=args.sheet)
    elif suffix == ".docx":
        raw = read_docx(args.input)
    elif suffix == ".pdf":
        raw = read_pdf(args.input)
    else:
        sys.exit(f"✗ 不支援的格式：{suffix}（支援 .xlsx/.xlsm/.csv/.docx/.pdf）")

    if not raw:
        sys.exit("✗ 沒有解析出任何列，請確認檔案結構")

    rows, conflicts = dedup(raw)
    n_codes = len({r["code"] for r in rows})
    multi = sum(1 for _, g in _group_by_code(rows).items() if len(g) > 1)

    print(f"\n原始 {len(raw)} 列 → 去重後 {len(rows)} 筆")
    print(f"相異 code {n_codes}｜一碼多變體的 code {multi} 個")
    if conflicts:
        print(f"⚠ {conflicts} 組同鍵但內容不一致，已取較完整者——建議抽樣人工核對")
    if multi:
        print(f"⚠ 本來源的 code 不唯一，匯入前必須確認 alarms 主鍵已含 variant，"
              f"否則 {len(rows)} 筆會被壓成 {n_codes} 筆")

    for f in ("cause", "action"):
        n = sum(1 for r in rows if r.get(f))
        print(f"  {f:8} 有值 {n}/{len(rows)} ({n * 100 // max(len(rows), 1)}%)")

    use_variant, variant_reason = decide_variant_mode(rows, args.variant_mode)
    print(f"\n[{args.variant_mode}] {variant_reason}", file=sys.stderr)
    if not use_variant and multi:
        print(f"\n✗ --variant-mode never 但來源有 {multi} 個代碼對應多筆，"
              f"{len(rows)} 筆會被壓成 {n_codes} 筆。請改用 always。")
        sys.exit(1)

    if args.against:
        warns = compare_against(rows, args.against)
        if warns:
            print(f"\n⚠ 措辭變動偵測：{len(warns)} 筆")
            for w in warns[:10]:
                print(f"  [WARN] {w}")
            if len(warns) > 10:
                print(f"  ...另外 {len(warns) - 10} 筆")
            print("  建議：確認是否為同一筆；若是，改用既有措辭或手動遷移 local_solution")

    issues = quality_check(rows)
    if issues:
        print(f"\n品質檢查：{len(issues)} 項")
        for level in ("ERROR", "WARN", "INFO"):
            sub = [i for i in issues if i[1] == level]
            for key, _, msg in sub[:8]:
                print(f"  [{level}] {key}: {msg}")
            if len(sub) > 8:
                print(f"  [{level}] ...另外 {len(sub) - 8} 項")
    else:
        print("\n品質檢查：無異常")

    if args.dry_run:
        print("\n--dry-run，未寫檔")
        return

    out = to_output(rows, args.device_model, args.action_to,
                    use_variant, args.description_format)
    dest = args.output or args.input.with_suffix(".json")
    dest.write_text(json.dumps(out, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"\n✓ 已寫入 {dest}（{len(out)} 筆，處置寫進 {args.action_to}，"
          f"variant {'啟用' if use_variant else '留空'}，description {args.description_format}）")
    print("⚠ 匯入前務必人工抽樣核對——這是製藥設備的處置指示，現場人員會照做。")


def _group_by_code(rows: list[dict]) -> dict:
    g: dict[str, list] = defaultdict(list)
    for r in rows:
        g[r["code"]].append(r)
    return g


if __name__ == "__main__":
    main()
